import streamlit as st
import json
import re
from datetime import datetime

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Resume & Portfolio Builder",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Serif+Display:ital@0;1&family=DM+Mono:wght@400;500&display=swap');

:root {
    --bg: #0D0D0D;
    --surface: #161616;
    --surface2: #1E1E1E;
    --accent: #C8FF00;
    --accent2: #FF6B35;
    --text: #F0EDE6;
    --muted: #6B6B6B;
    --border: #2A2A2A;
}

html, body, [class*="css"] {
    font-family: 'Syne', sans-serif;
    background-color: var(--bg) !important;
    color: var(--text) !important;
}

.stApp { background-color: var(--bg) !important; }

/* Sidebar */
[data-testid="stSidebar"] {
    background-color: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] .stRadio label {
    font-family: 'Syne', sans-serif !important;
    color: var(--text) !important;
    font-size: 14px !important;
}

/* Headers */
h1, h2, h3 { font-family: 'Syne', sans-serif !important; }

/* Inputs */
.stTextInput input, .stTextArea textarea, .stSelectbox select {
    background-color: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    color: var(--text) !important;
    border-radius: 4px !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 13px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px rgba(200,255,0,0.15) !important;
}

/* Buttons */
.stButton button {
    background-color: var(--accent) !important;
    color: #0D0D0D !important;
    border: none !important;
    border-radius: 2px !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    font-size: 13px !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    padding: 0.5rem 1.5rem !important;
    transition: all 0.2s ease !important;
}
.stButton button:hover {
    background-color: #D4FF1A !important;
    transform: translateY(-1px) !important;
}

/* Output boxes */
.output-box {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-left: 3px solid var(--accent);
    border-radius: 4px;
    padding: 1.5rem;
    font-family: 'DM Mono', monospace;
    font-size: 13px;
    line-height: 1.8;
    white-space: pre-wrap;
    color: var(--text);
    margin-top: 1rem;
}

/* Score badge */
.score-badge {
    display: inline-block;
    background: var(--accent);
    color: #0D0D0D;
    font-family: 'Syne', sans-serif;
    font-weight: 800;
    font-size: 2.5rem;
    padding: 0.5rem 1.5rem;
    border-radius: 2px;
}

.score-label {
    font-family: 'DM Mono', monospace;
    font-size: 11px;
    letter-spacing: 0.1em;
    color: var(--muted);
    text-transform: uppercase;
    margin-top: 4px;
}

/* Hero banner */
.hero {
    background: linear-gradient(135deg, #161616 0%, #1a1a0a 100%);
    border: 1px solid var(--border);
    border-top: 3px solid var(--accent);
    padding: 2.5rem 2rem;
    margin-bottom: 2rem;
    position: relative;
    overflow: hidden;
}
.hero::before {
    content: '✦';
    position: absolute;
    right: 2rem;
    top: 50%;
    transform: translateY(-50%);
    font-size: 8rem;
    color: rgba(200,255,0,0.04);
    font-family: 'Syne', sans-serif;
}
.hero-title {
    font-family: 'DM Serif Display', serif;
    font-size: 2.8rem;
    line-height: 1.1;
    color: var(--text);
    margin: 0;
}
.hero-title span { color: var(--accent); }
.hero-sub {
    font-family: 'DM Mono', monospace;
    font-size: 12px;
    letter-spacing: 0.1em;
    color: var(--muted);
    text-transform: uppercase;
    margin-top: 0.5rem;
}

/* Tag chips */
.tag {
    display: inline-block;
    background: var(--surface2);
    border: 1px solid var(--border);
    color: var(--accent);
    font-family: 'DM Mono', monospace;
    font-size: 10px;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    padding: 3px 10px;
    border-radius: 2px;
    margin: 2px;
}

/* Progress bar */
.progress-bar-container {
    background: var(--surface2);
    border-radius: 2px;
    height: 6px;
    margin: 8px 0;
}
.progress-bar-fill {
    height: 6px;
    border-radius: 2px;
    background: var(--accent);
    transition: width 0.5s ease;
}

/* Expanders */
.streamlit-expanderHeader {
    background-color: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    font-family: 'Syne', sans-serif !important;
    color: var(--text) !important;
}

/* Labels */
label { color: var(--muted) !important; font-size: 12px !important; letter-spacing: 0.05em !important; text-transform: uppercase !important; }

/* Dividers */
hr { border-color: var(--border) !important; }

/* Info boxes */
.stInfo { background-color: var(--surface2) !important; border-color: var(--accent) !important; }

/* Multiselect */
.stMultiSelect span { background-color: var(--surface2) !important; color: var(--accent) !important; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
#  HuggingFace LLM Wrapper  (uses free Inference API / can swap to local)
# ══════════════════════════════════════════════════════════════════════════════
def call_hf_api(prompt: str, max_tokens: int = 900) -> str:
    """
    Calls Groq API (free tier) — fast and reliable.
    Model: llama3-8b-8192
    """
    import requests, os
    api_key = st.secrets.get("GROQ_API_KEY", os.environ.get("GROQ_API_KEY", ""))

    if not api_key:
        return "[ERROR: No GROQ_API_KEY found. Please add it to Streamlit Secrets.]"

    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama3-8b-8192",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": 0.7,
        }
        r = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"[Groq API Error: {e}]"


def generate_pdf(text: str, title: str = "Resume") -> bytes:
    """Generate a clean PDF from text."""
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_margins(20, 20, 20)
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(4)
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("##") or line.isupper():
                pdf.set_font("Helvetica", "B", 11)
                pdf.ln(3)
                pdf.cell(0, 7, line.replace("##","").strip(), ln=True)
                pdf.set_font("Helvetica", "", 10)
            elif line == "":
                pdf.ln(3)
            else:
                pdf.multi_cell(0, 6, line)
        return bytes(pdf.output())
    except Exception as e:
        return b""


def generate_docx(text: str, title: str = "Resume") -> bytes:
    """Generate a Word document from text."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        t = doc.add_heading(title, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("##") or (line.isupper() and len(line) > 3):
                doc.add_heading(line.replace("##","").strip(), level=2)
            elif line == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(10)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        return b""


# ══════════════════════════════════════════════════════════════════════════════
#  ML Job Match Scorer  (keyword / TF-IDF cosine – no API needed)
# ══════════════════════════════════════════════════════════════════════════════
def compute_job_match(resume_text: str, job_description: str) -> dict:
    """
    Simple ML-based scorer using sklearn TF-IDF cosine similarity.
    Falls back to keyword overlap if sklearn not available.
    """
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        import numpy as np

        docs = [resume_text.lower(), job_description.lower()]
        vec = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
        tfidf = vec.fit_transform(docs)
        score = float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])
        score_pct = min(int(score * 180), 100)  # scale to 0-100

        # Extract missing keywords
        job_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', job_description.lower()))
        resume_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', resume_text.lower()))
        missing = list(job_words - resume_words)[:12]

        return {"score": score_pct, "missing_keywords": missing, "method": "TF-IDF Cosine Similarity"}
    except ImportError:
        # Fallback: keyword overlap
        job_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', job_description.lower()))
        resume_words = set(re.findall(r'\b[a-zA-Z]{4,}\b', resume_text.lower()))
        if not job_words:
            return {"score": 0, "missing_keywords": [], "method": "Keyword Overlap"}
        overlap = len(job_words & resume_words) / len(job_words)
        score_pct = min(int(overlap * 100), 100)
        missing = list(job_words - resume_words)[:12]
        return {"score": score_pct, "missing_keywords": missing, "method": "Keyword Overlap (sklearn not found)"}


# ══════════════════════════════════════════════════════════════════════════════
#  Prompt Builders
# ══════════════════════════════════════════════════════════════════════════════
def build_resume_prompt(data: dict) -> str:
    return f"""You are a professional resume writer. Create a polished, ATS-optimized resume in plain text format.

Student Profile:
- Name: {data['name']}
- Email: {data['email']} | Phone: {data['phone']}
- LinkedIn: {data.get('linkedin', 'N/A')} | GitHub: {data.get('github', 'N/A')}
- Degree: {data['degree']} in {data['major']} at {data['university']} ({data['grad_year']})
- GPA: {data.get('gpa', 'N/A')}
- Skills: {data['skills']}
- Projects: {data['projects']}
- Experience: {data.get('experience', 'N/A')}
- Certifications: {data.get('certifications', 'N/A')}
- Target Role: {data['target_role']}

Write a complete resume with sections: Summary, Education, Skills, Projects, Experience (if any), Certifications.
Make the Summary 2-3 sentences tailored to the target role. Be specific and quantify achievements where possible.
Output ONLY the resume text, no extra commentary."""


def build_cover_letter_prompt(data: dict, job_desc: str) -> str:
    return f"""You are an expert career coach. Write a compelling, personalized cover letter for a student applying to a job.

Student: {data['name']}
Degree: {data['degree']} in {data['major']} – {data['university']}
Skills: {data['skills']}
Projects: {data['projects']}
Target Role: {data['target_role']}

Job Description:
{job_desc}

Write a 3-paragraph cover letter (Opening hook, Body with 2-3 specific skill-project matches, Strong closing CTA).
Tone: Confident, professional, genuine. Avoid clichés like "I am writing to express my interest."
Output ONLY the cover letter text."""


def build_portfolio_prompt(data: dict) -> str:
    return f"""You are a web portfolio copywriter. Write compelling portfolio bio and project descriptions.

Student: {data['name']}
Major: {data['major']} | Graduation: {data['grad_year']}
Skills: {data['skills']}
Projects: {data['projects']}
Target Role: {data['target_role']}
Fun fact / Personal statement: {data.get('bio_note', 'Passionate about technology and solving real-world problems')}

Write the following sections:
1. HERO TAGLINE (one punchy line, max 12 words)
2. ABOUT ME (3 sentences, first-person, shows personality + ambition)
3. PROJECT CARDS (for each project mentioned, write: Title | One-line description | Key tech used | Impact/outcome)
4. SKILLS SECTION (categorize skills into groups like Languages, Frameworks, Tools, Soft Skills)

Output ONLY these sections with clear headings."""


# ══════════════════════════════════════════════════════════════════════════════
#  Sidebar Navigation
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div style='padding: 1rem 0; border-bottom: 1px solid #2A2A2A; margin-bottom: 1.5rem;'>
        <div style='font-family: DM Serif Display, serif; font-size: 1.4rem; color: #F0EDE6;'>✦ CareerForge</div>
        <div style='font-family: DM Mono, monospace; font-size: 10px; color: #6B6B6B; letter-spacing: 0.1em; text-transform: uppercase; margin-top: 4px;'>AI Resume & Portfolio MVP</div>
    </div>
    """, unsafe_allow_html=True)

    page = st.radio(
        "NAVIGATE",
        ["🏠  Home", "📄  Resume Generator", "✉️  Cover Letter", "🌐  Portfolio Builder", "🎯  Job Match Scorer"],
        label_visibility="visible"
    )

    st.markdown("---")
    st.markdown("""
    <div style='font-family: DM Mono, monospace; font-size: 10px; color: #6B6B6B;'>
    <div style='color: #C8FF00; margin-bottom: 6px;'>▸ STACK</div>
    Streamlit · HuggingFace<br>
    Mistral-7B · sklearn<br><br>
    <div style='color: #C8FF00; margin-bottom: 6px;'>▸ MVP VERSION</div>
    v0.1.0 · Feb 2025
    </div>
    """, unsafe_allow_html=True)

    with st.expander("⚙️ HF API Token"):
        hf_token = st.text_input("HuggingFace Token", type="password", placeholder="hf_...", help="Get free token at huggingface.co/settings/tokens")
        if hf_token:
            import os; os.environ["HF_TOKEN"] = hf_token
            st.success("Token set ✓")


# ══════════════════════════════════════════════════════════════════════════════
#  Session State – shared student data
# ══════════════════════════════════════════════════════════════════════════════
if "student_data" not in st.session_state:
    st.session_state.student_data = {}
if "resume_output" not in st.session_state:
    st.session_state.resume_output = ""


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE: HOME
# ══════════════════════════════════════════════════════════════════════════════
if "Home" in page:
    st.markdown("""
    <div class='hero'>
        <div class='hero-sub'>✦ Powered by HuggingFace · Mistral-7B</div>
        <div class='hero-title'>Build your career<br>story with <span>AI</span></div>
        <p style='font-family: DM Mono, monospace; font-size: 13px; color: #8A8A8A; margin-top: 1rem; max-width: 500px;'>
        From raw skills to a polished resume, tailored cover letter, and stunning portfolio — all in minutes.
        </p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    features = [
        ("📄", "Resume Generator", "ATS-optimized resumes tailored to your target role using Mistral-7B"),
        ("✉️", "Cover Letter", "Job-specific letters that match your skills to the job description"),
        ("🌐", "Portfolio Builder", "Hero copy, project cards, and bio generated for your web portfolio"),
        ("🎯", "Job Match Scorer", "ML-powered TF-IDF similarity score between your resume and any JD"),
    ]
    for col, (icon, title, desc) in zip([col1, col2, col3, col4], features):
        with col:
            st.markdown(f"""
            <div style='background:#161616; border:1px solid #2A2A2A; border-top:2px solid #C8FF00; padding:1.2rem; border-radius:2px; height:160px;'>
                <div style='font-size:1.5rem;'>{icon}</div>
                <div style='font-family:Syne,sans-serif; font-weight:700; font-size:0.85rem; margin: 8px 0 6px; color:#F0EDE6;'>{title}</div>
                <div style='font-family:DM Mono,monospace; font-size:11px; color:#6B6B6B; line-height:1.6;'>{desc}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("### 🚀 Roadmap — AI Depth Layers")
    roadmap = [
        ("MVP (Now)", "GenAI (Mistral-7B)", "Resume · Cover Letter · Portfolio · ML Job Scorer", "#C8FF00"),
        ("v0.2", "Fine-tuned LLM", "Domain-specific resume model trained on LinkedIn data", "#FF6B35"),
        ("v0.3", "Agentic AI", "Auto-scrape job boards, match & apply autonomously", "#8B5CF6"),
        ("v0.4", "DL / NLP", "BERT-based skill extraction, semantic job matching", "#06B6D4"),
        ("v1.0", "Full Platform", "Auth, history, PDF export, GitHub sync, interview prep", "#F59E0B"),
    ]
    for version, tech, desc, color in roadmap:
        st.markdown(f"""
        <div style='display:flex; align-items:center; gap:1rem; background:#161616; border:1px solid #2A2A2A; border-left:3px solid {color}; padding:0.8rem 1rem; margin-bottom:6px; border-radius:2px;'>
            <div style='font-family:DM Mono,monospace; font-size:11px; color:{color}; min-width:50px;'>{version}</div>
            <div style='font-family:Syne,sans-serif; font-weight:600; font-size:12px; color:#F0EDE6; min-width:150px;'>{tech}</div>
            <div style='font-family:DM Mono,monospace; font-size:11px; color:#6B6B6B;'>{desc}</div>
        </div>
        """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
#  SHARED STUDENT FORM (used by Resume + Cover Letter + Portfolio)
# ══════════════════════════════════════════════════════════════════════════════
def render_student_form(prefix=""):
    sd = st.session_state.student_data
    with st.expander("👤 Student Profile", expanded=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            name = st.text_input("Full Name", value=sd.get("name", ""), key=f"{prefix}name")
            email = st.text_input("Email", value=sd.get("email", ""), key=f"{prefix}email")
            phone = st.text_input("Phone", value=sd.get("phone", ""), key=f"{prefix}phone")
        with c2:
            degree = st.selectbox("Degree", ["B.Tech", "B.Sc", "BCA", "MCA", "M.Tech", "MBA", "Other"], key=f"{prefix}degree", index=["B.Tech","B.Sc","BCA","MCA","M.Tech","MBA","Other"].index(sd.get("degree","B.Tech")))
            major = st.text_input("Major / Branch", value=sd.get("major", "Computer Science"), key=f"{prefix}major")
            university = st.text_input("University", value=sd.get("university", ""), key=f"{prefix}university")
        with c3:
            grad_year = st.text_input("Graduation Year", value=sd.get("grad_year", "2025"), key=f"{prefix}grad_year")
            gpa = st.text_input("GPA / CGPA", value=sd.get("gpa", ""), key=f"{prefix}gpa")
            target_role = st.text_input("Target Role", value=sd.get("target_role", "Software Engineer"), key=f"{prefix}target_role", placeholder="e.g. ML Engineer, Data Analyst")

        c4, c5 = st.columns(2)
        with c4:
            linkedin = st.text_input("LinkedIn URL", value=sd.get("linkedin", ""), key=f"{prefix}linkedin")
            github = st.text_input("GitHub URL", value=sd.get("github", ""), key=f"{prefix}github")
        with c5:
            certifications = st.text_input("Certifications", value=sd.get("certifications", ""), key=f"{prefix}certifications", placeholder="e.g. AWS Cloud Practitioner, TensorFlow Developer")
            experience = st.text_area("Internship / Work Experience", value=sd.get("experience", ""), key=f"{prefix}experience", height=68, placeholder="Company · Role · Duration · Key achievement")

        skills = st.text_area("Technical Skills", value=sd.get("skills", ""), key=f"{prefix}skills", height=68, placeholder="Python, TensorFlow, React, SQL, Git, Docker...")
        projects = st.text_area("Projects (one per line)", value=sd.get("projects", ""), key=f"{prefix}projects", height=100, placeholder="Project Name – brief description – tech stack used\nAnother Project – ...")

    data = {
        "name": name, "email": email, "phone": phone, "degree": degree,
        "major": major, "university": university, "grad_year": grad_year,
        "gpa": gpa, "target_role": target_role, "linkedin": linkedin,
        "github": github, "certifications": certifications,
        "experience": experience, "skills": skills, "projects": projects,
    }
    st.session_state.student_data = data
    return data


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE: RESUME GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
if "Resume" in page:
    st.markdown("## 📄 Resume Generator")
    st.markdown("<div style='font-family:DM Mono,monospace;font-size:12px;color:#6B6B6B;margin-bottom:1.5rem;'>Powered by Mistral-7B via HuggingFace Inference API</div>", unsafe_allow_html=True)

    data = render_student_form("res_")

    col_btn, col_tip = st.columns([1, 3])
    with col_btn:
        generate = st.button("⚡ Generate Resume")
    with col_tip:
        st.markdown("<div style='font-family:DM Mono,monospace;font-size:11px;color:#6B6B6B;padding-top:10px;'>First run may take 20-30s (model warm-up)</div>", unsafe_allow_html=True)

    if generate:
        if not data["name"] or not data["skills"]:
            st.warning("Please fill in at least Name and Skills.")
        else:
            with st.spinner("Generating your resume with Mistral-7B..."):
                prompt = build_resume_prompt(data)
                result = call_hf_api(prompt, max_tokens=1000)
                st.session_state.resume_output = result

    if st.session_state.resume_output:
        st.markdown("### ✦ Generated Resume")
        st.markdown(f"<div class='output-box'>{st.session_state.resume_output}</div>", unsafe_allow_html=True)
        fname = data.get('name','student').replace(' ','_')
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("⬇ Download PDF", data=generate_pdf(st.session_state.resume_output, "Resume"), file_name=f"resume_{fname}.pdf", mime="application/pdf")
        with c2:
            st.download_button("⬇ Download Word", data=generate_docx(st.session_state.resume_output, "Resume"), file_name=f"resume_{fname}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with c3:
            st.download_button("⬇ Download TXT", data=st.session_state.resume_output, file_name=f"resume_{fname}.txt", mime="text/plain")


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE: COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
if "Cover" in page:
    st.markdown("## ✉️ Cover Letter Generator")
    st.markdown("<div style='font-family:DM Mono,monospace;font-size:12px;color:#6B6B6B;margin-bottom:1.5rem;'>Tailored to the specific job description you provide</div>", unsafe_allow_html=True)

    data = render_student_form("cl_")

    st.markdown("### 📋 Job Description")
    job_desc = st.text_area("Paste the full job description here", height=180, placeholder="Copy-paste the job posting here. The AI will match your skills to the requirements...")

    col_btn, _ = st.columns([1, 3])
    with col_btn:
        generate_cl = st.button("⚡ Generate Cover Letter")

    if generate_cl:
        if not data["name"] or not job_desc:
            st.warning("Please fill in your Name and paste a Job Description.")
        else:
            with st.spinner("Crafting your personalized cover letter..."):
                prompt = build_cover_letter_prompt(data, job_desc)
                result = call_hf_api(prompt, max_tokens=700)

            st.markdown("### ✦ Generated Cover Letter")
            st.markdown(f"<div class='output-box'>{result}</div>", unsafe_allow_html=True)
            fname = data.get('name','student').replace(' ','_')
            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("⬇ Download PDF", data=generate_pdf(result, "Cover Letter"), file_name=f"cover_letter_{fname}.pdf", mime="application/pdf")
            with c2:
                st.download_button("⬇ Download Word", data=generate_docx(result, "Cover Letter"), file_name=f"cover_letter_{fname}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with c3:
                st.download_button("⬇ Download TXT", data=result, file_name=f"cover_letter_{fname}.txt", mime="text/plain")


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE: PORTFOLIO BUILDER
# ══════════════════════════════════════════════════════════════════════════════
if "Portfolio" in page:
    st.markdown("## 🌐 Portfolio Builder")
    st.markdown("<div style='font-family:DM Mono,monospace;font-size:12px;color:#6B6B6B;margin-bottom:1.5rem;'>Generates copy for your personal portfolio website</div>", unsafe_allow_html=True)

    data = render_student_form("pf_")

    bio_note = st.text_input("Personal Statement / Fun Fact", placeholder="e.g. I love building AI tools that help students land their first job")
    data["bio_note"] = bio_note

    col_btn, _ = st.columns([1, 3])
    with col_btn:
        generate_pf = st.button("⚡ Generate Portfolio Copy")

    if generate_pf:
        if not data["name"] or not data["projects"]:
            st.warning("Please fill in Name and at least one Project.")
        else:
            with st.spinner("Writing your portfolio content..."):
                prompt = build_portfolio_prompt(data)
                result = call_hf_api(prompt, max_tokens=900)

            st.markdown("### ✦ Portfolio Content")
            st.markdown(f"<div class='output-box'>{result}</div>", unsafe_allow_html=True)
            fname = data.get('name','student').replace(' ','_')
            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("⬇ Download PDF", data=generate_pdf(result, "Portfolio"), file_name=f"portfolio_{fname}.pdf", mime="application/pdf")
            with c2:
                st.download_button("⬇ Download Word", data=generate_docx(result, "Portfolio"), file_name=f"portfolio_{fname}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with c3:
                st.download_button("⬇ Download TXT", data=result, file_name=f"portfolio_{fname}.txt", mime="text/plain")


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE: JOB MATCH SCORER
# ══════════════════════════════════════════════════════════════════════════════
if "Job Match" in page:
    st.markdown("## 🎯 Job Match Scorer")
    st.markdown("<div style='font-family:DM Mono,monospace;font-size:12px;color:#6B6B6B;margin-bottom:1.5rem;'>ML-powered TF-IDF Cosine Similarity — No API needed</div>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Your Resume Text**")
        resume_text = st.text_area(
            "Paste your resume",
            value=st.session_state.resume_output,
            height=300,
            label_visibility="collapsed",
            placeholder="Paste your full resume text here (or generate one from the Resume tab first)..."
        )
    with col2:
        st.markdown("**Job Description**")
        jd_text = st.text_area(
            "Paste job description",
            height=300,
            label_visibility="collapsed",
            placeholder="Paste the job description here..."
        )

    col_btn, _ = st.columns([1, 3])
    with col_btn:
        score_btn = st.button("⚡ Calculate Match Score")

    if score_btn:
        if not resume_text or not jd_text:
            st.warning("Please provide both resume text and job description.")
        else:
            with st.spinner("Calculating similarity..."):
                result = compute_job_match(resume_text, jd_text)

            score = result["score"]
            missing = result["missing_keywords"]
            method = result["method"]

            st.markdown("---")
            c1, c2, c3 = st.columns([1, 1, 2])
            with c1:
                color = "#C8FF00" if score >= 70 else "#FF6B35" if score >= 40 else "#FF3333"
                st.markdown(f"""
                <div style='text-align:center; padding: 1rem;'>
                    <div style='font-family:DM Serif Display,serif; font-size:4rem; color:{color}; line-height:1;'>{score}%</div>
                    <div style='font-family:DM Mono,monospace; font-size:10px; color:#6B6B6B; text-transform:uppercase; letter-spacing:0.1em; margin-top:6px;'>Match Score</div>
                </div>
                """, unsafe_allow_html=True)
            with c2:
                label = "Strong Match ✓" if score >= 70 else "Moderate Match" if score >= 40 else "Weak Match ✗"
                tips = "Your resume aligns well with this role." if score >= 70 else "Add more relevant keywords from the JD." if score >= 40 else "Significant gaps found. Tailor your resume."
                st.markdown(f"""
                <div style='background:#161616; border:1px solid #2A2A2A; border-left:3px solid {color}; padding:1rem; border-radius:2px;'>
                    <div style='font-family:Syne,sans-serif; font-weight:700; color:{color}; font-size:14px;'>{label}</div>
                    <div style='font-family:DM Mono,monospace; font-size:11px; color:#8A8A8A; margin-top:6px; line-height:1.6;'>{tips}</div>
                    <div style='font-family:DM Mono,monospace; font-size:9px; color:#4A4A4A; margin-top:8px;'>Algorithm: {method}</div>
                </div>
                """, unsafe_allow_html=True)
            with c3:
                if missing:
                    st.markdown("**Missing Keywords from JD**")
                    chips = "".join([f"<span class='tag'>{kw}</span>" for kw in missing[:12]])
                    st.markdown(f"<div style='margin-top:6px;'>{chips}</div>", unsafe_allow_html=True)
                    st.markdown(f"<div style='font-family:DM Mono,monospace;font-size:11px;color:#6B6B6B;margin-top:10px;'>Add these to your resume to improve ATS score</div>", unsafe_allow_html=True)
                else:
                    st.markdown("<div style='font-family:DM Mono,monospace;font-size:12px;color:#C8FF00;'>✓ No significant keyword gaps found!</div>", unsafe_allow_html=True)

            # Visual progress bar
            st.markdown(f"""
            <div style='margin-top:1.5rem;'>
                <div style='font-family:DM Mono,monospace;font-size:10px;color:#6B6B6B;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:6px;'>Match Strength</div>
                <div class='progress-bar-container'>
                    <div class='progress-bar-fill' style='width:{score}%; background:{color};'></div>
                </div>
                <div style='display:flex;justify-content:space-between;font-family:DM Mono,monospace;font-size:9px;color:#4A4A4A;margin-top:4px;'>
                    <span>0%</span><span>50%</span><span>100%</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
